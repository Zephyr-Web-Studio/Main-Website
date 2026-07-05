import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def generate_brochure():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Set page background color to dark #050505 (deep black / slate)
    background_xml = f'<w:background {nsdecls("w")} w:color="050505"/>'
    doc.element.insert(0, parse_xml(background_xml))
    
    settings_xml = f'<w:displayBackgroundShape {nsdecls("w")}/>'
    doc.settings.element.append(parse_xml(settings_xml))
    
    # Define Color Palette
    COLOR_WHITE = RGBColor(255, 255, 255)
    COLOR_LIME = RGBColor(163, 230, 53)      # #a3e635 (Neon Green)
    COLOR_GRAY = RGBColor(163, 163, 163)      # #a3a3a3 (Muted Gray)
    COLOR_DARK_GRAY = RGBColor(100, 100, 100) # Dark gray for subtext
    
    HEX_DARK_BG = "050505"
    HEX_CARD_BG = "121214"                    # Dark slate card background
    HEX_ZEBRA_BG = "0B0B0C"                   # Alternate table row background
    HEX_LIME = "A3E635"
    HEX_BORDER = "242427"
    
    # Helper: Style Run
    def style_run(run, font_name="Outfit", font_size=11, bold=False, italic=False, color=COLOR_WHITE):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        
    # Helper: Style Paragraph
    def style_paragraph(p, space_before=0, space_after=6, line_spacing=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        
    # Helper: Cell background and padding
    def style_cell(cell, fill_hex=None, top_padding=120, bottom_padding=120, left_padding=150, right_padding=150):
        tcPr = cell._tc.get_or_add_tcPr()
        if fill_hex:
            existing_shd = tcPr.find(qn('w:shd'))
            if existing_shd is not None:
                tcPr.remove(existing_shd)
            shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
            tcPr.append(parse_xml(shading_xml))
            
        # Set cell padding in dxa (1 pt = 20 dxa)
        tcMar = OxmlElement('w:tcMar')
        for margin_name, val in [('top', top_padding), ('bottom', bottom_padding), ('left', left_padding), ('right', right_padding)]:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
        
    # Helper: Cell Borders
    def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
        tcPr = cell._tc.get_or_add_tcPr()
        existing = tcPr.find(qn('w:tcBorders'))
        if existing is not None:
            tcPr.remove(existing)
            
        tcBorders = OxmlElement('w:tcBorders')
        borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
        for edge, data in borders.items():
            node = OxmlElement(f'w:{edge}')
            if data:
                node.set(qn('w:val'), data.get('val', 'single'))
                node.set(qn('w:sz'), str(data.get('sz', 4)))
                node.set(qn('w:space'), '0')
                node.set(qn('w:color'), data.get('color', 'auto'))
            else:
                node.set(qn('w:val'), 'none')
            tcBorders.append(node)
        tcPr.append(tcBorders)

    # Helper: Bullet Point
    def add_bullet_point(doc, text, bullet_char="•", color=COLOR_WHITE):
        p = doc.add_paragraph()
        style_paragraph(p, space_before=2, space_after=4)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        
        r_bullet = p.add_run(f"{bullet_char}\t")
        style_run(r_bullet, font_name="Outfit", font_size=10.5, bold=True, color=COLOR_LIME)
        
        r_text = p.add_run(text)
        style_run(r_text, font_name="Outfit", font_size=10.5, color=color)
        return p

    # Helper: Numbered Point
    def add_numbered_point(doc, num, text):
        p = doc.add_paragraph()
        style_paragraph(p, space_before=2, space_after=4)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        
        r_num = p.add_run(f"{num}.\t")
        style_run(r_num, font_name="Outfit", font_size=10.5, bold=True, color=COLOR_LIME)
        
        r_text = p.add_run(text)
        style_run(r_text, font_name="Outfit", font_size=10.5, color=COLOR_WHITE)
        return p

    # Helper: Colored Separator Line
    def add_separator_line(doc, color_hex, height_pt=2, space_after=12):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        style_cell(cell, fill_hex=color_hex, top_padding=0, bottom_padding=0, left_padding=0, right_padding=0)
        
        trPr = table.rows[0]._tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_pt * 20)))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        style_paragraph(p, space_before=0, space_after=space_after)
        
    # Helper: Set Column Widths
    def set_col_widths(table, widths):
        for row in table.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = w

    # Configure Header/Footer (no header/footer on cover)
    doc.sections[0].different_first_page_header_footer = True
    header = doc.sections[0].header
    footer = doc.sections[0].footer
    
    p_h = header.paragraphs[0]
    style_paragraph(p_h, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r_h = p_h.add_run("ZEPHYR WEB STUDIO  |  COMPANY BROCHURE")
    style_run(r_h, font_name="Outfit", font_size=8.5, color=COLOR_GRAY)
    
    p_f = footer.paragraphs[0]
    style_paragraph(p_f)
    r_f_left = p_f.add_run("Zephyr Web Studio Brochure")
    style_run(r_f_left, font_name="Outfit", font_size=8.5, color=COLOR_GRAY)
    p_f.add_run("\t\t")
    r_f_right = p_f.add_run("Page ")
    style_run(r_f_right, font_name="Outfit", font_size=8.5, color=COLOR_GRAY)
    
    # Custom Page Numbering Helper
    def add_page_number(run):
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
    add_page_number(r_f_right)

    # --- PAGE 1: COVER PAGE ---
    p_spacer_top = doc.add_paragraph()
    style_paragraph(p_spacer_top, space_after=60)
    
    # Brand Logo
    p_logo = doc.add_paragraph()
    style_paragraph(p_logo, space_after=36, align=WD_ALIGN_PARAGRAPH.CENTER)
    logo_path = 'public/Zephyr.png'
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(2.2))
    else:
        r_logo = p_logo.add_run("ZEPHYR WEB STUDIO")
        style_run(r_logo, font_name="Syne", font_size=28, bold=True, color=COLOR_LIME)
        
    # Main Headline
    p_head = doc.add_paragraph()
    style_paragraph(p_head, space_after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_head_w1 = p_head.add_run("Looking to Upgrade\nYour ")
    style_run(r_head_w1, font_name="Syne", font_size=32, bold=True, color=COLOR_WHITE)
    r_head_w2 = p_head.add_run("Website?")
    style_run(r_head_w2, font_name="Syne", font_size=32, bold=True, color=COLOR_LIME)
    
    p_spacer_mid = doc.add_paragraph()
    style_paragraph(p_spacer_mid, space_after=24)
    
    # Pill Button (Consultation CTA)
    btn_table = doc.add_table(rows=1, cols=1)
    btn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    btn_cell = btn_table.cell(0, 0)
    style_cell(btn_cell, fill_hex=HEX_LIME, top_padding=160, bottom_padding=160, left_padding=300, right_padding=300)
    set_cell_borders(btn_cell)
    
    p_btn = btn_cell.paragraphs[0]
    style_paragraph(p_btn, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_btn = p_btn.add_run("Get Free Consultation Today!")
    style_run(r_btn, font_name="Outfit", font_size=12.5, bold=True, color=RGBColor(0, 0, 0))
    set_col_widths(btn_table, [Inches(3.6)])
    
    doc.add_page_break()

    # --- PAGE 2: OUR SERVICES ---
    p_srv_h = doc.add_paragraph()
    style_paragraph(p_srv_h, space_before=12, space_after=12)
    r_srv_h = p_srv_h.add_run("OUR SERVICES")
    style_run(r_srv_h, font_name="Syne", font_size=22, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=24)
    
    # 3-row vertical Services list
    srv_table = doc.add_table(rows=3, cols=1)
    srv_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    services = [
        ("🎨 Web Design", "Bespoke user experiences (UX), custom themes, interactive mockups, and mobile-responsive layout guides tailored to your brand identity."),
        ("💻 Web Development", "High-performance websites, portals, and dashboards engineered with React, Next.js, and TypeScript, backed by cloud hosting and databases."),
        ("📱 App Development", "High-velocity cross-platform mobile apps for iOS & Android, equipped with live tracking, video calling, and secure payment systems.")
    ]
    
    for idx, (title, desc) in enumerate(services):
        cell = srv_table.cell(idx, 0)
        
        style_cell(cell, fill_hex=HEX_CARD_BG, top_padding=160, bottom_padding=160, left_padding=240, right_padding=240)
        
        # Left-side lime green vertical accent border on each service card
        set_cell_borders(cell, 
                         top={"val": "single", "sz": 2, "color": HEX_BORDER},
                         bottom={"val": "single", "sz": 2, "color": HEX_BORDER},
                         right={"val": "single", "sz": 2, "color": HEX_BORDER},
                         left={"val": "single", "sz": 12, "color": HEX_LIME})
                         
        p_title = cell.paragraphs[0]
        style_paragraph(p_title, space_after=6)
        r_title = p_title.add_run(title)
        style_run(r_title, font_name="Syne", font_size=13, bold=True, color=COLOR_WHITE)
        
        p_desc = cell.add_paragraph()
        style_paragraph(p_desc, space_after=0)
        r_desc = p_desc.add_run(desc)
        style_run(r_desc, font_name="Outfit", font_size=9.5, color=COLOR_GRAY)
        
    set_col_widths(srv_table, [Inches(6.4)])
    
    p_srv_note = doc.add_paragraph()
    style_paragraph(p_srv_note, space_before=40, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_srv_note = p_srv_note.add_run("Work with us  •  www.zephyrwebstudio.dev")
    style_run(r_srv_note, font_name="Outfit", font_size=10, bold=True, color=COLOR_LIME)

    doc.add_page_break()

    # --- PAGE 3: ABOUT US ---
    p_abt_h = doc.add_paragraph()
    style_paragraph(p_abt_h, space_before=12, space_after=12)
    r_abt_h = p_abt_h.add_run("ABOUT US")
    style_run(r_abt_h, font_name="Syne", font_size=22, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=36)
    
    # Large quote-like callout card for About Us content
    abt_table = doc.add_table(rows=1, cols=1)
    abt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    abt_cell = abt_table.cell(0, 0)
    
    style_cell(abt_cell, fill_hex=HEX_CARD_BG, top_padding=300, bottom_padding=300, left_padding=300, right_padding=300)
    set_cell_borders(abt_cell, 
                     top={"val": "single", "sz": 4, "color": HEX_BORDER},
                     bottom={"val": "single", "sz": 4, "color": HEX_BORDER},
                     right={"val": "single", "sz": 4, "color": HEX_BORDER},
                     left={"val": "single", "sz": 16, "color": HEX_LIME})
                     
    p_abt = abt_cell.paragraphs[0]
    style_paragraph(p_abt, space_after=0, line_spacing=1.3)
    r_abt = p_abt.add_run(
        "We are a website service provider that pushes the boundaries of traditional web development. "
        "Our edgy approach to web design sets us apart, and our team of talented developers is dedicated "
        "to bringing your vision to life. Get in touch. Let us help you make an impact and leave a lasting impression."
    )
    style_run(r_abt, font_name="Outfit", font_size=12, color=COLOR_WHITE)
    
    set_col_widths(abt_table, [Inches(6.4)])

    doc.add_page_break()

    # --- PAGE 4: WHY CHOOSE US ---
    p_why_h = doc.add_paragraph()
    style_paragraph(p_why_h, space_before=12, space_after=12)
    r_why_h = p_why_h.add_run("WHY CHOOSE US")
    style_run(r_why_h, font_name="Syne", font_size=22, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=24)
    
    # 2x2 grid of stats cards (3 cards standard size, 1 highlighted large card)
    why_table = doc.add_table(rows=2, cols=2)
    why_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    stats = [
        ("2+", "years of industry experience", HEX_CARD_BG, COLOR_LIME),
        ("40+", "brand partners", HEX_CARD_BG, COLOR_LIME),
        ("15+", "team members", HEX_CARD_BG, COLOR_LIME),
        ("100%", "Customer Satisfaction", HEX_LIME, RGBColor(0, 0, 0)) # Highlight card: Lime green bg, black text
    ]
    
    for idx, (value, desc, bg_hex, text_color) in enumerate(stats):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = why_table.cell(row_idx, col_idx)
        
        style_cell(cell, fill_hex=bg_hex, top_padding=240, bottom_padding=240, left_padding=200, right_padding=200)
        
        if bg_hex == HEX_LIME:
            set_cell_borders(cell)
        else:
            set_cell_borders(cell, 
                             top={"val": "single", "sz": 2, "color": HEX_BORDER},
                             bottom={"val": "single", "sz": 2, "color": HEX_BORDER},
                             left={"val": "single", "sz": 2, "color": HEX_BORDER},
                             right={"val": "single", "sz": 2, "color": HEX_BORDER})
                             
        p_val = cell.paragraphs[0]
        style_paragraph(p_val, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        r_val = p_val.add_run(value)
        style_run(r_val, font_name="Syne", font_size=28, bold=True, color=text_color)
        
        p_desc = cell.add_paragraph()
        style_paragraph(p_desc, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        r_desc = p_desc.add_run(desc)
        style_run(r_desc, font_name="Outfit", font_size=10, bold=True, color=text_color if bg_hex == HEX_LIME else COLOR_GRAY)
        
    set_col_widths(why_table, [Inches(3.2), Inches(3.2)])

    doc.add_page_break()

    # --- PAGE 5: OUR FEATURES ---
    p_feat_h = doc.add_paragraph()
    style_paragraph(p_feat_h, space_before=12, space_after=12)
    r_feat_h = p_feat_h.add_run("OUR FEATURES")
    style_run(r_feat_h, font_name="Syne", font_size=22, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=18)
    
    # 2 columns layout: Left is Core Capabilities (Dark gray card), Right is Advanced Solutions (Dark gray card)
    feat_table = doc.add_table(rows=1, cols=2)
    feat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left_feat = feat_table.cell(0, 0)
    c_right_feat = feat_table.cell(0, 1)
    
    # Left column: Core Capabilities (Dark gray card with lime left border)
    style_cell(c_left_feat, fill_hex=HEX_CARD_BG, top_padding=200, bottom_padding=200, left_padding=240, right_padding=240)
    set_cell_borders(c_left_feat, 
                     top={"val": "single", "sz": 4, "color": HEX_BORDER},
                     bottom={"val": "single", "sz": 4, "color": HEX_BORDER},
                     left={"val": "single", "sz": 12, "color": HEX_LIME},
                     right={"val": "single", "sz": 2, "color": HEX_BORDER})
    
    p_bf_h = c_left_feat.paragraphs[0]
    style_paragraph(p_bf_h, space_after=10)
    r_bf_h = p_bf_h.add_run("CORE CAPABILITIES")
    style_run(r_bf_h, font_name="Syne", font_size=12, bold=True, color=COLOR_LIME)
    
    core_feats = [
        ("UI/UX Design", "Custom brand templates, interfaces, and user journey mockups."),
        ("SEO Optimization", "Semantic markup structures engineered to rank highly on search engines."),
        ("Mobile Optimization", "Fluid responsive layouts for seamless browsing on phones and tablets."),
        ("Interactive Forms", "Custom inputs, customer enquiry flows, and contact portals.")
    ]
    for title, desc in core_feats:
        p_title = c_left_feat.add_paragraph()
        style_paragraph(p_title, space_before=4, space_after=2)
        r_bullet = p_title.add_run("•  ")
        style_run(r_bullet, font_name="Outfit", font_size=10, bold=True, color=COLOR_LIME)
        r_t = p_title.add_run(title)
        style_run(r_t, font_name="Outfit", font_size=10, bold=True, color=COLOR_WHITE)
        
        p_d = c_left_feat.add_paragraph()
        style_paragraph(p_d, space_before=0, space_after=6)
        p_d.paragraph_format.left_indent = Inches(0.18)
        r_d = p_d.add_run(desc)
        style_run(r_d, font_name="Outfit", font_size=8.5, color=COLOR_GRAY)
        
    # Right column: Advanced Solutions (Dark gray card with lime left border)
    style_cell(c_right_feat, fill_hex=HEX_CARD_BG, top_padding=200, bottom_padding=200, left_padding=240, right_padding=240)
    set_cell_borders(c_right_feat, 
                     top={"val": "single", "sz": 4, "color": HEX_BORDER},
                     bottom={"val": "single", "sz": 4, "color": HEX_BORDER},
                     left={"val": "single", "sz": 12, "color": HEX_LIME},
                     right={"val": "single", "sz": 4, "color": HEX_BORDER})
                     
    p_tf_h = c_right_feat.paragraphs[0]
    style_paragraph(p_tf_h, space_after=10)
    r_tf_h = p_tf_h.add_run("ADVANCED SOLUTIONS")
    style_run(r_tf_h, font_name="Syne", font_size=12, bold=True, color=COLOR_LIME)
    
    adv_feats = [
        ("Voice & Video Call", "Integrated real-time audio and video communications in-app."),
        ("Media Galleries", "Premium banner slides, lightboxes, and inline video integrations."),
        ("AWS Integration", "Cloud servers setup, secure databases, and storage infrastructure."),
        ("1-Year Support", "Full developer support and corrections included for a whole year.")
    ]
    for title, desc in adv_feats:
        p_title = c_right_feat.add_paragraph()
        style_paragraph(p_title, space_before=4, space_after=2)
        r_bullet = p_title.add_run("•  ")
        style_run(r_bullet, font_name="Outfit", font_size=10, bold=True, color=COLOR_LIME)
        r_t = p_title.add_run(title)
        style_run(r_t, font_name="Outfit", font_size=10, bold=True, color=COLOR_WHITE)
        
        p_d = c_right_feat.add_paragraph()
        style_paragraph(p_d, space_before=0, space_after=6)
        p_d.paragraph_format.left_indent = Inches(0.18)
        r_d = p_d.add_run(desc)
        style_run(r_d, font_name="Outfit", font_size=8.5, color=COLOR_GRAY)
        
    set_col_widths(feat_table, [Inches(3.2), Inches(3.2)])

    doc.add_page_break()

    # --- PAGE 6: WEB DESIGNING PROCESS ---
    p_prc_h = doc.add_paragraph()
    style_paragraph(p_prc_h, space_before=12, space_after=12)
    r_prc_h = p_prc_h.add_run("WEB DESIGNING PROCESS")
    style_run(r_prc_h, font_name="Syne", font_size=20, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=18)
    
    # Process steps vertical layout using standard tables
    steps_table = doc.add_table(rows=5, cols=2)
    steps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    steps = [
        ("01", "Advance Payment", "The process begins with a 50% advance payment from the client for website design."),
        ("02", "Content Submission", "You need to send all content related to your business."),
        ("03", "Website Work", "The website work will begin according to your business requirements and will be completed in 4-5 working days."),
        ("04", "Final Payment & Website Live", "After the remaining balance is paid, the website is live and amendments can be made."),
        ("05", "Corrections", "Website corrections will be done by the company as per your need for one whole year.")
    ]
    
    for row_idx, (num, title, desc) in enumerate(steps):
        cell_num = steps_table.cell(row_idx, 0)
        cell_desc = steps_table.cell(row_idx, 1)
        
        bg_color = HEX_CARD_BG if row_idx % 2 == 0 else HEX_ZEBRA_BG
        style_cell(cell_num, fill_hex=bg_color, top_padding=120, bottom_padding=120, left_padding=150, right_padding=150)
        style_cell(cell_desc, fill_hex=bg_color, top_padding=120, bottom_padding=120, left_padding=150, right_padding=150)
        
        border_color = HEX_BORDER
        set_cell_borders(cell_num, 
                         top={"val": "single", "sz": 2, "color": border_color} if row_idx > 0 else None,
                         bottom={"val": "single", "sz": 2, "color": border_color} if row_idx < 4 else None,
                         left={"val": "single", "sz": 12, "color": HEX_LIME})
        set_cell_borders(cell_desc, 
                         top={"val": "single", "sz": 2, "color": border_color} if row_idx > 0 else None,
                         bottom={"val": "single", "sz": 2, "color": border_color} if row_idx < 4 else None,
                         right={"val": "single", "sz": 2, "color": border_color})
                         
        p_num = cell_num.paragraphs[0]
        style_paragraph(p_num, space_after=0)
        r_num = p_num.add_run(num)
        style_run(r_num, font_name="Syne", font_size=18, bold=True, color=COLOR_LIME)
        
        p_desc = cell_desc.paragraphs[0]
        style_paragraph(p_desc, space_after=4)
        r_title = p_desc.add_run(f"{title}\n")
        style_run(r_title, font_name="Outfit", font_size=11, bold=True, color=COLOR_WHITE)
        
        r_desc = p_desc.add_run(desc)
        style_run(r_desc, font_name="Outfit", font_size=9.5, color=COLOR_GRAY)
        
    set_col_widths(steps_table, [Inches(1.2), Inches(5.2)])

    doc.add_page_break()

    # --- PAGE 7: OUR PRICING (SILVER & GOLD) ---
    p_prc_h2 = doc.add_paragraph()
    style_paragraph(p_prc_h2, space_before=12, space_after=12)
    r_prc_h2 = p_prc_h2.add_run("OUR PRICING")
    style_run(r_prc_h2, font_name="Syne", font_size=20, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=18)
    
    # Columns side-by-side: Silver vs Gold Plan cards
    pricing_table = doc.add_table(rows=1, cols=2)
    pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_silver = pricing_table.cell(0, 0)
    c_gold = pricing_table.cell(0, 1)
    
    # Silver Plan Card (Dark background, white text)
    style_cell(c_silver, fill_hex=HEX_CARD_BG, top_padding=200, bottom_padding=200, left_padding=200, right_padding=200)
    set_cell_borders(c_silver, 
                     top={"val": "single", "sz": 4, "color": HEX_BORDER},
                     bottom={"val": "single", "sz": 4, "color": HEX_BORDER},
                     left={"val": "single", "sz": 4, "color": HEX_BORDER},
                     right={"val": "single", "sz": 2, "color": HEX_BORDER})
                     
    p_sil_h = c_silver.paragraphs[0]
    style_paragraph(p_sil_h, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_sil_h = p_sil_h.add_run("Silver Plan")
    style_run(r_sil_h, font_name="Syne", font_size=14, bold=True, color=COLOR_GRAY)
    
    silver_features = [
        "Single-page responsive website",
        "School intro & contact info",
        "Basic photo gallery",
        "Contact/Inquiry form",
        "Quick updates for announcements",
        "News/Notice scrolling section",
        "1 round of text/image revision"
    ]
    for feat in silver_features:
        p = c_silver.add_paragraph()
        style_paragraph(p, space_before=2, space_after=2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        r_bullet = p.add_run("- ")
        style_run(r_bullet, font_name="Outfit", font_size=9, color=COLOR_GRAY)
        r_txt = p.add_run(feat)
        style_run(r_txt, font_name="Outfit", font_size=9.5, color=COLOR_WHITE)
        
    # Gold Plan Card (Lime Green background, black text)
    style_cell(c_gold, fill_hex=HEX_LIME, top_padding=200, bottom_padding=200, left_padding=200, right_padding=200)
    set_cell_borders(c_gold)
    
    p_gld_h = c_gold.paragraphs[0]
    style_paragraph(p_gld_h, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_gld_h = p_gld_h.add_run("Gold Plan")
    style_run(r_gld_h, font_name="Syne", font_size=14, bold=True, color=RGBColor(0, 0, 0))
    
    gold_features = [
        "Up to 5 pages",
        "Custom school branding/colors",
        "Staff/faculty showcase",
        "Events calendar",
        "Document download section",
        "News & Updates section",
        "Announcement pop-ups",
        "Photo or event gallery",
        "Google Maps, FAQ",
        "Blog/Noticeboard area",
        "2 rounds of revisions"
    ]
    for feat in gold_features:
        p = c_gold.add_paragraph()
        style_paragraph(p, space_before=2, space_after=2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        r_bullet = p.add_run("- ")
        style_run(r_bullet, font_name="Outfit", font_size=9, color=RGBColor(0, 0, 0))
        r_txt = p.add_run(feat)
        style_run(r_txt, font_name="Outfit", font_size=9.5, color=RGBColor(0, 0, 0))
        
    set_col_widths(pricing_table, [Inches(3.2), Inches(3.2)])
    
    p_prc_note = doc.add_paragraph()
    style_paragraph(p_prc_note, space_before=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_prc_note1 = p_prc_note.add_run("The listed prices do not include the cost of the domain name.\n")
    style_run(r_prc_note1, font_name="Outfit", font_size=9, color=COLOR_GRAY)
    r_prc_note2 = p_prc_note.add_run("Every extra page will cost ₹300 for each page")
    style_run(r_prc_note2, font_name="Outfit", font_size=10, bold=True, color=COLOR_LIME)

    doc.add_page_break()

    # --- PAGE 8: OUR PRICING (PLATINUM PLAN) ---
    p_prc_h3 = doc.add_paragraph()
    style_paragraph(p_prc_h3, space_before=12, space_after=12)
    r_prc_h3 = p_prc_h3.add_run("OUR PRICING (CONTINUED)")
    style_run(r_prc_h3, font_name="Syne", font_size=20, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=18)
    
    # Large featured card for Platinum Plan
    plat_table = doc.add_table(rows=1, cols=1)
    plat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_plat = plat_table.cell(0, 0)
    
    style_cell(c_plat, fill_hex=HEX_LIME, top_padding=240, bottom_padding=240, left_padding=240, right_padding=240)
    set_cell_borders(c_plat)
    
    # Premium Badge
    p_plat_badge = c_plat.paragraphs[0]
    style_paragraph(p_plat_badge, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_plat_badge = p_plat_badge.add_run("★  PREMIUM  ★")
    style_run(r_plat_badge, font_name="Outfit", font_size=9, bold=True, color=RGBColor(255, 215, 0)) # Gold Color text
    
    p_plat_h = c_plat.add_paragraph()
    style_paragraph(p_plat_h, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_plat_h = p_plat_h.add_run("Platinum Plan")
    style_run(r_plat_h, font_name="Syne", font_size=16, bold=True, color=RGBColor(0, 0, 0))
    
    plat_features = [
        "Up to 10 custom pages",
        "All Gold features included",
        "Student portal",
        "Parent resources/downloads",
        "Achievements and alumni showcase",
        "Class timetable section",
        "Integrated calendar with events/holidays",
        "Enquiry/Admission form with custom fields",
        "Social media integration",
        "Staff directory with photos & contacts",
        "On-page SEO optimization",
        "Testimonials & feedback section",
        "Priority support",
        "3 rounds of revisions"
    ]
    for feat in plat_features:
        p = c_plat.add_paragraph()
        style_paragraph(p, space_before=2, space_after=2)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        r_bullet = p.add_run("✔\t") # Clean checkmark bullet
        style_run(r_bullet, font_name="Outfit", font_size=9, bold=True, color=RGBColor(0, 0, 0))
        r_txt = p.add_run(feat)
        style_run(r_txt, font_name="Outfit", font_size=10, color=RGBColor(0, 0, 0))
        
    set_col_widths(plat_table, [Inches(6.4)])
    
    p_plat_note = doc.add_paragraph()
    style_paragraph(p_plat_note, space_before=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_plat_note1 = p_plat_note.add_run("The listed prices do not include the cost of the domain name.\n")
    style_run(r_plat_note1, font_name="Outfit", font_size=9, color=COLOR_GRAY)
    r_plat_note2 = p_plat_note.add_run("Every extra page will cost ₹300 for each page")
    style_run(r_plat_note2, font_name="Outfit", font_size=10, bold=True, color=COLOR_LIME)

    doc.add_page_break()

    # --- PAGE 9: CONTACT US ---
    p_cnt_h = doc.add_paragraph()
    style_paragraph(p_cnt_h, space_before=12, space_after=12)
    r_cnt_h = p_cnt_h.add_run("CONTACT US")
    style_run(r_cnt_h, font_name="Syne", font_size=22, bold=True, color=COLOR_LIME)
    add_separator_line(doc, HEX_LIME, height_pt=2, space_after=24)
    
    # Visual card for contact details inside mobile mock-like container
    contact_table = doc.add_table(rows=1, cols=1)
    contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_cell = contact_table.cell(0, 0)
    
    style_cell(c_cell, fill_hex=HEX_CARD_BG, top_padding=300, bottom_padding=300, left_padding=240, right_padding=240)
    set_cell_borders(c_cell, 
                     top={"val": "single", "sz": 12, "color": HEX_LIME},
                     bottom={"val": "single", "sz": 4, "color": HEX_BORDER},
                     left={"val": "single", "sz": 4, "color": HEX_BORDER},
                     right={"val": "single", "sz": 4, "color": HEX_BORDER})
                     
    # Phone numbers row
    p_ph = c_cell.paragraphs[0]
    style_paragraph(p_ph, space_after=12)
    r_ph_icon = p_ph.add_run("📞\t")
    style_run(r_ph_icon, font_name="Outfit", font_size=12, color=COLOR_LIME)
    r_ph_num1 = p_ph.add_run("+91 85498 40238\n")
    style_run(r_ph_num1, font_name="Outfit", font_size=12, bold=True, color=COLOR_WHITE)
    r_ph_num2 = p_ph.add_run("\t+91 72093 33695\n")
    style_run(r_ph_num2, font_name="Outfit", font_size=12, bold=True, color=COLOR_WHITE)
    r_ph_num3 = p_ph.add_run("\t+91 7050 218002")
    style_run(r_ph_num3, font_name="Outfit", font_size=12, bold=True, color=COLOR_WHITE)
    
    # Email row
    p_email = c_cell.add_paragraph()
    style_paragraph(p_email, space_before=6, space_after=12)
    r_email_icon = p_email.add_run("✉\t")
    style_run(r_email_icon, font_name="Outfit", font_size=12, color=COLOR_LIME)
    r_email_txt = p_email.add_run("forwork.zephyrwebstudio@gmail.com")
    style_run(r_email_txt, font_name="Outfit", font_size=12, color=COLOR_WHITE)
    
    # Instagram row
    p_ig = c_cell.add_paragraph()
    style_paragraph(p_ig, space_before=6, space_after=12)
    r_ig_icon = p_ig.add_run("📷\t")
    style_run(r_ig_icon, font_name="Outfit", font_size=12, color=COLOR_LIME)
    r_ig_txt = p_ig.add_run("@zephyrwebstudio (Instagram)")
    style_run(r_ig_txt, font_name="Outfit", font_size=11, color=COLOR_WHITE)
    
    # X (Twitter) row
    p_tw = c_cell.add_paragraph()
    style_paragraph(p_tw, space_before=6, space_after=12)
    r_tw_icon = p_tw.add_run("𝕏\t")
    style_run(r_tw_icon, font_name="Outfit", font_size=12, color=COLOR_LIME)
    r_tw_txt = p_tw.add_run("@zephyrwebstudio (X / Twitter)")
    style_run(r_tw_txt, font_name="Outfit", font_size=11, color=COLOR_WHITE)
    
    # YouTube row
    p_yt = c_cell.add_paragraph()
    style_paragraph(p_yt, space_before=6, space_after=12)
    r_yt_icon = p_yt.add_run("📺\t")
    style_run(r_yt_icon, font_name="Outfit", font_size=12, color=COLOR_LIME)
    r_yt_txt = p_yt.add_run("@zephyrwebstudio (YouTube)")
    style_run(r_yt_txt, font_name="Outfit", font_size=11, color=COLOR_WHITE)
    
    # Location row
    p_loc = c_cell.add_paragraph()
    style_paragraph(p_loc, space_before=6, space_after=18)
    r_loc_icon = p_loc.add_run("📍\t")
    style_run(r_loc_icon, font_name="Outfit", font_size=12, color=COLOR_LIME)
    r_loc_txt = p_loc.add_run("Bangalore, India")
    style_run(r_loc_txt, font_name="Outfit", font_size=12, color=COLOR_WHITE)
    
    # Spacer
    add_separator_line(c_cell, HEX_BORDER, height_pt=1, space_after=12)
    
    # Support status
    p_sup = c_cell.add_paragraph()
    style_paragraph(p_sup, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_sup = p_sup.add_run("⚡  24 x 7 Support  ⚡")
    style_run(r_sup, font_name="Syne", font_size=11, bold=True, color=COLOR_LIME)
    
    set_col_widths(contact_table, [Inches(6.4)])

    # Save to disk
    output_filename = "Zephyr_Brochure_English.docx"
    doc.save(output_filename)
    print(f"Brochure document saved successfully as: {output_filename}")

if __name__ == "__main__":
    generate_brochure()
